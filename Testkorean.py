import os
import json
import time
import pytz
from datetime import datetime
from docx import Document
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import shutil

from openai import AzureOpenAI
from dotenv import load_dotenv

# ==============================
# CONFIG
# ==============================
URL = "https://englishdart.fss.or.kr/dsbh001/main.do?rcpNo=20260310901403"
load_dotenv()

# ==============================
# ENV
# ==============================
def get_env(key):
    try:
        import streamlit as st
        if key in st.secrets:
            return st.secrets[key]
    except:
        pass
    return os.getenv(key)

# ==============================
# OPENAI CLIENT
# ==============================
def get_client():
    return AzureOpenAI(
        api_key=get_env("AZURE_OPENAI_API_KEY"),
        api_version=get_env("AZURE_OPENAI_API_VERSION"),
        azure_endpoint=get_env("AZURE_OPENAI_ENDPOINT")
    )

# ==============================
# SELENIUM DRIVER (CLOUD SAFE)
# ==============================
def get_driver():

    options = Options()
    options.add_argument("--headless")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--disable-gpu")
    options.add_argument("--window-size=1920,1080")
    options.add_argument("--disable-blink-features=AutomationControlled")

    # Detect browser
    chromium_path = (
        shutil.which("chromium") or
        shutil.which("chromium-browser") or
        shutil.which("google-chrome") or
        shutil.which("google-chrome-stable")
    )
    if chromium_path:
        options.binary_location = chromium_path

    chromedriver_path = (
        shutil.which("chromedriver") or
        shutil.which("chromium-driver")
    )

    if chromedriver_path:
        driver = webdriver.Chrome(service=Service(chromedriver_path), options=options)
    else:
        driver = webdriver.Chrome(options=options)

    return driver

# ==============================
# SCRAPE KOREAN DART
# ==============================
def scrape_korean_dart(driver, wait, url):

    print("Opening Korean DART homepage...")
    driver.get("https://englishdart.fss.or.kr")
    wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    time.sleep(2)

    print("Opening disclosure page...")
    driver.get(url)

    time.sleep(5)

    # Try iframe
    try:
        iframe = wait.until(EC.presence_of_element_located((By.TAG_NAME, "iframe")))
        driver.switch_to.frame(iframe)
        print("Switched to iframe")
    except:
        print("No iframe, continuing...")

    time.sleep(2)

    # Scroll
    driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
    time.sleep(2)

    text = driver.find_element(By.TAG_NAME, "body").text

    print("Extracted text length:", len(text))

    return text

# ==============================
# OPENAI EXTRACTION
# ==============================
def extract_data(text):

    client = get_client()

    prompt = f"""
Read the disclosure text and return:

1️⃣ Summary in English
2️⃣ Extract these parameters

Return JSON only. Do not add explanation.

Parameters to extract:

Project / Asset Name
Country / City
Sector / Sub-sector
Investment / Deal Value
Currency
Sponsors / Investors
Lenders / Banks
Project Status / Stage
Source / Publication Date

If a value is missing write "Not Found".

TEXT:
{text[:15000]}

Output JSON format:

{{
"summary":"English summary here",
"parameters":{{
"Project / Asset Name":"",
"Country / City":"",
"Sector / Sub-sector":"",
"Investment / Deal Value":"",
"Currency":"",
"Sponsors / Investors":"",
"Lenders / Banks":"",
"Project Status / Stage":"",
"Source / Publication Date":""
}}
}}
"""

    response = client.chat.completions.create(
        model=get_env("AZURE_OPENAI_DEPLOYMENT"),
        messages=[
            {"role": "system", "content": "You extract structured financial project data and always return valid JSON only."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.2
    )

    result = response.choices[0].message.content
    return json.loads(result)

    

# ==============================
# WORD EXPORT
# ==============================
def create_word(data):

    summary = data["summary"]
    parameters = data["parameters"]

    #onedrive_path = os.getenv("OneDriveCommercial") or os.getenv("OneDrive")
    # ✅ LOCAL OUTPUT FOLDER
    output_folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
    os.makedirs(output_folder, exist_ok=True)



    tz = pytz.timezone("Asia/Kolkata")
    now_local = datetime.now(tz)
    timestamp = now_local.strftime("%Y-%m-%d_%H-%M-%S")

    file_name = f"DART_Summary_{timestamp}.docx"

    file_path = os.path.join(output_folder, file_name)

    doc = Document()

    doc.add_heading("Korean DART Disclosure Summary", level=0)

    doc.add_paragraph(f"Generated on: {now_local.strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("Source URL", level=1)
    doc.add_paragraph(URL)

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)

    doc.add_paragraph("")

    # TABLE
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    headers = table.rows[0].cells
    headers[0].text = "Parameters"
    headers[1].text = "Value"

    for cell in headers:
        run = cell.paragraphs[0].runs[0]
        run.bold = True

    for key, value in parameters.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value

    doc.save(file_path)

    # JSON metadata
    json_file_path = file_path.replace(".docx", ".json")

    metadata = {
        "file_name": file_name,
        "project_name": parameters.get("Project / Asset Name", "Not Found"),
        "country": "South Korea",
        "region": "APAC",
        "industry_type": parameters.get("Sector / Sub-sector", "Not Found"),
        "generated_date": now_local.strftime("%Y-%m-%d %H:%M:%S"),
        "file_size_kb": round(os.path.getsize(file_path) / 1024, 1),
        "source_url": URL,
        "website": "Korea Dart"
    }

    with open(json_file_path, "w") as json_file:
        json.dump(metadata, json_file, indent=2)

    print("Word file saved:", file_path)
    print("Metadata saved to:", json_file_path)

# ==============================
# MAIN
# ==============================

def run():

    print("Extracting website text...")

    driver = get_driver()
    wait = WebDriverWait(driver, 60)

    text = scrape_korean_dart(driver, wait, URL)

    print("Sending to Azure OpenAI...")

    data = extract_data(text)

    print("\nSummary:\n", data["summary"])

    print("\nCreating Word file...")

    create_word(data)

if __name__ == "__main__":
    run()
