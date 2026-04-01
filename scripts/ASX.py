import json
import pprint
import pytz
from datetime import datetime
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
import shutil
import time
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re

from NewsIntentOpenAI import detect_news_intent
from ParametersExtract import extract_project_details


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def accept_cookies(driver, wait):
    try:
        btn = wait.until(EC.element_to_be_clickable((By.XPATH, "//button[contains(text(),'Accept All Cookies')]")))
        btn.click()
        print("Cookies accepted.")
        time.sleep(2)
    except:
        print("No cookie popup found.")


def agree_and_proceed(driver, wait):
    try:
        # Try multiple possible selectors
        btn = None
        selectors = [
            (By.XPATH, "//*[contains(text(),'Agree and proceed')]"),
            (By.XPATH, "//*[contains(text(),'Agree and Proceed')]"),
            (By.XPATH, "//*[contains(@value,'Agree')]"),
            (By.XPATH, "//*[contains(@class,'agree')]"),
            (By.CSS_SELECTOR, "input[value*='Agree']"),
            (By.CSS_SELECTOR, "button.agree"),
        ]
        for by, selector in selectors:
            try:
                btn = WebDriverWait(driver, 5).until(EC.element_to_be_clickable((by, selector)))
                if btn:
                    print(f"Found agree button with: {selector}")
                    break
            except:
                continue

        if btn:
            btn.click()
            print("Agreed and proceeded.")
            time.sleep(3)
        else:
            # Debug: print page source snippet to find correct selector
            print("Agree button not found. Page snippet:")
            print(driver.page_source[:2000])
    except Exception as e:
        print(f"Agree error: {e}")


def run_asx_scraper():

    print("Starting ASX Energy scraper...")

    all_projects_text = []

    options = Options()
    options.add_argument("--headless=new")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--disable-gpu")
    options.add_argument("--window-size=1920,1080")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_experimental_option("excludeSwitches", ["enable-automation"])
    options.add_experimental_option("useAutomationExtension", False)

    driver = webdriver.Chrome(options=options)
    wait = WebDriverWait(driver, 30)

    try:
        url = "https://www.asx.com.au/markets/trade-our-cash-market/todays-announcements"
        print(f"\nOpening: {url}")

        driver.get(url)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(3)

        # Step 1: Accept cookies
        accept_cookies(driver, wait)

        # Scroll to trigger lazy load
        driver.execute_script("window.scrollTo(0, 600);")
        time.sleep(5)

        # Wait for table - check main page first, then iframe
        try:
            wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "tbody tr")))
        except:
            print("Checking inside iframes...")
            iframes = driver.find_elements(By.TAG_NAME, "iframe")
            for iframe in iframes:
                try:
                    driver.switch_to.frame(iframe)
                    if driver.find_elements(By.CSS_SELECTOR, "tbody tr"):
                        print("Found table inside iframe")
                        break
                    driver.switch_to.default_content()
                except:
                    driver.switch_to.default_content()

        rows = driver.find_elements(By.CSS_SELECTOR, "tbody tr")
        print(f"Total announcements found: {len(rows)}")

        energy_links = []

        for row in rows:
            try:
                cells = row.find_elements(By.TAG_NAME, "td")
                if not cells:
                    continue
                for cell in cells:
                    try:
                        link_elem = cell.find_element(By.TAG_NAME, "a")
                        raw_html = link_elem.get_attribute("innerHTML")
                        headline_plain = re.sub(r'<[^>]+>', '', raw_html)
                        headline_plain = re.sub(r'\s+', ' ', headline_plain).strip()
                        if "energy" in headline_plain.lower():
                            pdf_url = link_elem.get_attribute("href")
                            if not pdf_url.startswith("http"):
                                pdf_url = "https://www.asx.com.au" + pdf_url
                            asx_code = cells[0].text.strip()
                            date_time = cells[1].text.strip() if len(cells) > 1 else ""
                            # Clean headline - keep only first line (actual title)
                            headline_clean = headline_plain.split('\n')[0].strip()
                            energy_links.append({
                                "asx_code": asx_code,
                                "date_time": date_time,
                                "headline": headline_clean,
                                "pdf_url": pdf_url
                            })
                            print(f"Found Energy announcement: {headline_clean}")
                            break
                    except:
                        continue
            except:
                continue

        print(f"\nTotal Energy announcements: {len(energy_links)}")

        main_window = driver.current_window_handle

        for item in energy_links:
            print(f"\nProcessing: {item['headline']}")

            # Open PDF URL in new tab
            driver.execute_script(f"window.open('{item['pdf_url']}', '_blank');")
            time.sleep(3)

            new_tab = [h for h in driver.window_handles if h != main_window][0]
            driver.switch_to.window(new_tab)
            time.sleep(2)

            # Step 2: Click "Agree and proceed" if shown
            agree_and_proceed(driver, wait)
            time.sleep(3)

            # After agree, get the actual redirected PDF URL and navigate as text
            current_url = driver.current_url
            print(f"Current URL after agree: {current_url}")

            # Navigate to PDF URL with ?display=html or extract via embed
            # Use Chrome PDF viewer workaround: navigate to about:blank then use fetch
            pdf_text = ""
            try:
                # Try getting text from embed/iframe inside PDF viewer
                embeds = driver.find_elements(By.TAG_NAME, "embed")
                if embeds:
                    driver.switch_to.frame(embeds[0])
                    pdf_text = driver.find_element(By.TAG_NAME, "body").text.strip()
                    driver.switch_to.default_content()
            except:
                pass

            if not pdf_text:
                # Use JavaScript to fetch PDF page as text via Google PDF viewer trick
                try:
                    driver.get(f"https://docs.google.com/viewer?url={current_url}&embedded=true")
                    time.sleep(5)
                    pdf_text = driver.find_element(By.TAG_NAME, "body").text.strip()
                except:
                    pass

            if not pdf_text:
                # Last resort: get page source and extract visible text
                pdf_text = driver.execute_script("return document.body.innerText")
                pdf_text = pdf_text.strip() if pdf_text else ""

            print(f"DEBUG pdf_text length: {len(pdf_text)}")
            print(f"DEBUG pdf_text preview: {pdf_text[:300]}")


            # Clean PDF viewer UI noise
            pdf_text = re.sub(r'PDF\s*-*>?\s*\d+\s*page\s*[\d.]+KB', '', pdf_text, flags=re.IGNORECASE)
            pdf_text = re.sub(r'\d+\s*page\s*[\d.]+KB', '', pdf_text, flags=re.IGNORECASE)
            pdf_text = re.sub(r'Access to this site.*?I confirm that any content I access will not be used for any commercial purpose[^.]*\.', '', pdf_text, flags=re.DOTALL)
            pdf_text = pdf_text.strip()

            driver.close()
            driver.switch_to.window(main_window)

            if not pdf_text:
                print("No text extracted, skipping...")
                continue

            all_projects_text.append({
                "asx_code": item['asx_code'],
                "headline": item['headline'],
                "date_time": item['date_time'],
                "text": pdf_text,
                "source_url": item['pdf_url']
            })

    finally:
        driver.quit()

    print("\n===== FULL ARRAY DATA =====\n")
    pprint.pprint(all_projects_text)

    print("\nStarting OpenAI Processing...\n")

    for project in all_projects_text:

        print("Processing:", project["headline"])

        intent = detect_news_intent(project["text"])
        print("AI Intent:", intent)

        structured_data = extract_project_details(project["text"])

        print("\nExtracted JSON:\n")
        print(structured_data)

        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "output")
        os.makedirs(output_dir, exist_ok=True)

        if isinstance(structured_data, str):
            try:
                structured_data = json.loads(structured_data)
            except:
                structured_data = {}

        if project.get("date_time"):
            structured_data["Publication Date"] = project["date_time"]

        if project.get("source_url"):
            structured_data["Source Link"] = project["source_url"]

        project_name = structured_data.get("Project Name", project["headline"])
        project_summary = structured_data.get("Project Summary", "No summary available.")

        tz = pytz.timezone("Asia/Kolkata")
        now_local = datetime.now(tz)
        timestamp = now_local.strftime("%Y-%m-%d_%H-%M-%S")

        asx_code_safe = project["asx_code"].replace("/", "_")
        project_name_safe = re.sub(r'[\\/*?:"<>|]', "", project_name).replace(" ", "")

        file_name = f"ASX_{asx_code_safe}_{project_name_safe}_{timestamp}.docx"
        file_path = os.path.join(output_dir, file_name)

        doc = Document()
        doc.add_heading(project_name, level=0)

        gen_time = now_local.strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Generated on: {gen_time}")

        doc.add_heading("Project Summary", level=1)
        doc.add_paragraph(project_summary)
        doc.add_paragraph("")

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Parameters"
        hdr_cells[1].text = "Value"

        for i in range(2):
            para = hdr_cells[i].paragraphs[0]
            if para.runs:
                para.runs[0].bold = True

        keys = [k for k in structured_data.keys() if k not in ["Project Name", "Project Summary", "Source Link"]]

        for key in keys:
            value = structured_data[key]
            row_cells = table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)

        doc.add_paragraph("")
        doc.add_heading("Signals that Make a News Item Relevant", level=1)

        if isinstance(intent, list) and len(intent) >= 2:
            if intent[0].strip().lower().startswith("detected intent"):
                doc.add_paragraph(f"1. {intent[0].strip()}: {intent[1].strip()}")
                for it in intent[2:]:
                    doc.add_paragraph(str(it), style="List Bullet")
            else:
                doc.add_paragraph(f"1. {' '.join(str(i) for i in intent)}")
        elif isinstance(intent, str):
            items = [i.strip() for i in re.split(r'\n|\r|\d+\.\s*|\*\s*|-\s*', intent) if i.strip()]
            if len(items) >= 2 and items[0].lower().startswith("detected intent"):
                doc.add_paragraph(f"1. {items[0]}: {items[1]}")
                for it in items[2:]:
                    doc.add_paragraph(it, style="List Bullet")
            elif len(items) > 1:
                doc.add_paragraph(f"1. {' '.join(str(i) for i in items)}")
            else:
                doc.add_paragraph(f"1. {intent.strip()}")
        else:
            doc.add_paragraph(f"1. {str(intent)}")

        doc.add_paragraph("")
        if "Source Link" in structured_data:
            doc.add_heading("Source Link", level=1)
            source_link_url = structured_data["Source Link"]
            p = doc.add_paragraph()
            add_hyperlink(p, source_link_url, source_link_url)

        doc.save(file_path)
        print(f"Saved Word document: {file_path}")

        json_file_name = file_name.replace(".docx", ".json")
        json_file_path = os.path.join(output_dir, json_file_name)

        file_size_kb = os.path.getsize(file_path) / 1024

        metadata = {
            "file_name": file_name,
            "project_name": project_name,
            "asx_code": project["asx_code"],
            "headline": project["headline"],
            "date_time": project["date_time"],
            "country": "Australia",
            "region": "Asia-Pacific",
            "industry_type": "Energy",
            "website": "ASX",
            "generated_date": gen_time,
            "file_size_kb": round(file_size_kb, 1),
            "source_url": project.get("source_url", "N/A")
        }

        with open(json_file_path, "w") as json_file:
            json.dump(metadata, json_file, indent=2)

        print(f"Saved metadata JSON: {json_file_path}")
        print("\n----------------------------------\n")


if __name__ == "__main__":
    run_asx_scraper()
