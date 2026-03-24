import json
import pprint
from datetime import datetime, timedelta
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
import time
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re
import platform
from selenium import webdriver
from selenium.webdriver.chrome.options import Options

 
from NewsIntentOpenAI import detect_news_intent
from ParametersExtract import extract_project_details


def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set hyperlink style (blue and underlined)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink
 
 
def run_scraper():
 
    print("Opening sites.json...")
 
    with open("config/sites.json", "r") as f:
        data = json.load(f)
 
    sites = data["sites"]
 
    all_projects_text = []
 
    options = Options()
    if platform.system() == "Linux":
       options.binary_location = "/usr/bin/chromium"
    options.add_argument("--headless")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--disable-gpu")
    options.add_argument("--window-size=1920,1080")
    options.add_argument("--disable-extensions")
    options.add_argument("--disable-software-rasterizer")

    

    driver = webdriver.Chrome(options=options)

    
    wait = WebDriverWait(driver, 60)

    try:
        for site in sites:
 
            if site["valid"] != "Yes":
                print("Skipping site:", site["siteName"])
                continue
 
            print("\nOpening site:", site["siteName"])
 
            driver.get(site["siteURL"])
            wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
            time.sleep(2)
 
            industries = site["industryType"]
 
            # Support string or list
            if isinstance(industries, str):
                industries = [industries]
 
            for industry in industries:
 
                print("\nSelecting Industry:", industry)

                driver.get(site["siteURL"])
                wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
                time.sleep(2)

                try:
                    dropdown = wait.until(EC.presence_of_element_located((By.ID, "dropdown_3")))
                    Select(dropdown).select_by_visible_text(industry)
                except:
                    print(f"Industry '{industry}' not found in dropdown. Skipping...")
                    continue
 
                days = site["NoOfDays"]
 
                date_from = (datetime.today() - timedelta(days=days)).strftime("%Y-%m-%d")
                date_to = datetime.today().strftime("%Y-%m-%d")
 
                print("Setting date range:", date_from, "to", date_to)
 
                for field_id, value in [("date_4", date_from), ("date_5", date_to)]:
                    driver.execute_script(
                        """const el = document.getElementById(arguments[0]);
                        el.value = arguments[1];
                        el.dispatchEvent(new Event('input', { bubbles: true }));
                        el.dispatchEvent(new Event('change', { bubbles: true }));""",
                        field_id, value
                    )

                time.sleep(1)
 
                print("Clicking Apply Filter")

                try:
                    driver.execute_script("""
                        var btn = document.querySelector('button.btn-entitylist-filter-submit');
                        var target = btn.getAttribute('data-target');
                        var entitylist = jQuery(btn).closest('.entitylist');
                        var grid = entitylist.find('.entity-grid').filter(':first');
                        var formData = jQuery(target).find('input,select').serialize();
                        grid.trigger('metafilter', formData);
                    """)
                except Exception as e:
                    print(f"Apply Filter error: {e}")
                    continue
 
                # Wait for page to process the filter
                time.sleep(3)
                # Check if results table exists
                try:
                    wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "table tbody tr")))
                except:
                    print(f"No results found for industry '{industry}' in date range {date_from} to {date_to}")
                    print("Skipping to next industry...")
                    continue
                
                # Check if there are actual data rows (not just "No records found" message)
                rows = driver.find_elements(By.CSS_SELECTOR, "table tbody tr")
                if len(rows) == 0:
                    print(f"No records found for industry '{industry}'")
                    continue
                
                # Check if it's an empty state message
                first_row_text = rows[0].text.lower() if rows else ""
                if "no records" in first_row_text or "no results" in first_row_text:
                    print(f"No records found for industry '{industry}'")
                    continue
 
                # Find column indexes
                header_cells = driver.find_elements(By.CSS_SELECTOR, "table thead tr th")
                valid_date_idx = None
                location_idx = None
 
                for idx, th in enumerate(header_cells):
                    header_text = th.text.strip().lower()
 
                    if "valid date" in header_text:
                        valid_date_idx = idx
 
                    if "location" in header_text:
                        location_idx = idx
 
                rows = driver.find_elements(By.CSS_SELECTOR, "table tbody tr")
 
                print("Total records found:", len(rows))
 
                for i in range(len(rows)):
 
                    rows = driver.find_elements(By.CSS_SELECTOR, "table tbody tr")
 
                    cells = rows[i].find_elements(By.TAG_NAME, "td")
                    try:
                        link = cells[0].find_element(By.TAG_NAME, "a")
                    except:
                        continue
 
                    epbc_number = link.text
 
                    # Extract table values
                    valid_date = None
                    location = None
 
                    if valid_date_idx is not None and len(cells) > valid_date_idx:
                        valid_date = cells[valid_date_idx].text.strip()
 
                    if location_idx is not None and len(cells) > location_idx:
                        location = cells[location_idx].text.strip()
 
                    print("\nOpening record:", epbc_number)
 
                    link.click()
 
                    try:
                        wait.until(EC.presence_of_element_located((By.XPATH, "//*[contains(text(),'Project description')]")))
                    except:
                        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
 
                    time.sleep(3)
 
                    driver.execute_script("window.scrollTo(0, document.body.scrollHeight)")
                    time.sleep(1.5)
 
                    # Capture the project page URL after navigation
                    source_url = driver.current_url
 
                    text = driver.find_element(By.TAG_NAME, "body").text
 
                    all_projects_text.append({
                        "epbc_number": epbc_number,
                        "text": text,
                        "valid_date": valid_date,
                        "location": location,
                        "source_url": source_url,
                        "site_info": {
                            "region": site.get("region", "N/A"),
                            "country": site.get("country", "N/A"),
                            "siteName": site.get("siteName", "N/A"),
                            "industry": industry
                        }
                    })
 
                    driver.back()
 
                    wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
                    time.sleep(2)
                    
                    try:
                        driver.execute_script("""
                            var btn = document.querySelector('button.btn-entitylist-filter-submit');
                            var target = btn.getAttribute('data-target');
                            var entitylist = jQuery(btn).closest('.entitylist');
                            var grid = entitylist.find('.entity-grid').filter(':first');
                            var formData = jQuery(target).find('input,select').serialize();
                            grid.trigger('metafilter', formData);
                        """)
                    except:
                        print("Warning: Could not trigger metafilter after going back")
 
                    # Wait for table to reload after going back
                    try:
                        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "table tbody tr")))
                    except:
                        print("Warning: Table did not reload after going back. Continuing...")
                        # Try to continue anyway

    finally:
        driver.quit()
 
    print("\n===== FULL ARRAY DATA =====\n")
    pprint.pprint(all_projects_text)
 
    print("\nStarting OpenAI Processing...\n")
 
    for project in all_projects_text:
 
        print("Processing:", project["epbc_number"])
 
        intent = detect_news_intent(project["text"])
        print("AI Intent:", intent)
 
        structured_data = extract_project_details(project["text"])
 
        print("\nExtracted JSON:\n")
        print(structured_data)
 
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "output")
 
        os.makedirs(output_dir, exist_ok=True)
 
        if isinstance(structured_data, str):
            try:
                structured_data = json.loads(structured_data)
            except:
                structured_data = {}
 
        if project.get("valid_date"):
            structured_data["Publication Date"] = project["valid_date"]
 
        if project.get("location"):
            structured_data["Location"] = project["location"]
 
        # Add Source Link (project page URL) to structured_data
        if project.get("source_url"):
            structured_data["Source Link"] = project["source_url"]
 
        project_name = structured_data.get("Project Name", "Unknown Project")
        project_summary = structured_data.get("Project Summary", "No summary available.")
 
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
 
        epbc_safe = project["epbc_number"].replace("/", "_")
 
        project_name_safe = re.sub(r'[\\/*?:"<>|]', "", project_name)
        project_name_safe = project_name_safe.replace(" ", "")
 
        file_name = f"{epbc_safe}_{project_name_safe}_{timestamp}.docx"
 
        file_path = os.path.join(output_dir, file_name)
 
        doc = Document()
 
        doc.add_heading(project_name, level=0)
 
        gen_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Generated on: {gen_time}")
 
        doc.add_heading("Project Summary", level=1)
        doc.add_paragraph(project_summary)
 
        doc.add_paragraph("")
 
 
 
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
 
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Parameters"
        hdr_cells[1].text = "Value"
        # Make header row bold
        for i in range(2):
            para = hdr_cells[i].paragraphs[0]
            if para.runs:
                para.runs[0].bold = True
            else:
                run = para.add_run(hdr_cells[i].text)
                run.bold = True
                hdr_cells[i].text = ""
                para._element.remove(para.runs[0]._element)
 
        # Prepare ordered list of keys: all except Project Name/Project Summary/Source Link
        keys = [k for k in structured_data.keys() if k not in ["Project Name", "Project Summary", "Source Link"]]
 
        for key in keys:
            value = structured_data[key]
            row_cells = table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)
 
 
        doc.add_paragraph("")
        doc.add_heading("Signals that Make a News Item Relevant", level=1)
 
        # Detected Intent formatting as numbered line
        detected_intent = None
        if isinstance(intent, list) and len(intent) >= 2:
            # If first item is 'Detected Intent' and second is the value
            if intent[0].strip().lower().startswith("detected intent"):
                detected_intent = f"1. {intent[0].strip()}: {intent[1].strip()}"
                doc.add_paragraph(detected_intent)
                # Add any remaining items as bullets
                for item in intent[2:]:
                    doc.add_paragraph(str(item), style="List Bullet")
            else:
                # Fallback: join all as one numbered line
                doc.add_paragraph(f"1. {' '.join(str(i) for i in intent)}")
        elif isinstance(intent, str):
            # Try to split by newlines or numbered/bulleted list if present
            items = [i.strip() for i in re.split(r'\n|\r|\d+\.\s*|\*\s*|-\s*', intent) if i.strip()]
            if len(items) >= 2 and items[0].lower().startswith("detected intent"):
                detected_intent = f"1. {items[0]}: {items[1]}"
                doc.add_paragraph(detected_intent)
                for item in items[2:]:
                    doc.add_paragraph(item, style="List Bullet")
            elif len(items) > 1:
               
                doc.add_paragraph(f"1. {' '.join(str(i) for i in items)}")
            else:
                doc.add_paragraph(f"1. {intent.strip()}")
        else:
            doc.add_paragraph(f"1. {str(intent)}")
        
        # Add Source Link as a separate heading section with clickable hyperlink (after Signals section)
        doc.add_paragraph("")
        if "Source Link" in structured_data:
            doc.add_heading("Source Link", level=1)
            source_link_url = structured_data["Source Link"]
            p = doc.add_paragraph()
            add_hyperlink(p, source_link_url, source_link_url)
 
        doc.save(file_path)
 
        print(f"Saved Word document: {file_path}")
        
        # Create metadata JSON file
        json_file_name = file_name.replace(".docx", ".json")
        json_file_path = os.path.join(output_dir, json_file_name)
        
        # Get file size
        file_size_kb = os.path.getsize(file_path) / 1024
        
        # Prepare metadata
        site_info = project.get("site_info", {})
        metadata = {
            "file_name": file_name,
            "project_name": project_name,
            "country": structured_data.get("Country", site_info.get("country", "N/A")),
            "region": site_info.get("region", "N/A"),
            "industry_type": site_info.get("industry", "N/A"),
            "website": site_info.get("siteName", "N/A"),
            "generated_date": gen_time,
            "file_size_kb": round(file_size_kb, 1),
            "source_url": project.get("source_url", "N/A")
        }
        
        # Save metadata JSON
        with open(json_file_path, "w") as json_file:
            json.dump(metadata, json_file, indent=2)
        
        print(f"Saved metadata JSON: {json_file_path}")
        print("\n----------------------------------\n")
 
 
if __name__ == "__main__":
    run_scraper()
 