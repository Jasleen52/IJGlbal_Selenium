import sys
import os
import subprocess
import streamlit as st
import json
import glob
import time
import zipfile
from datetime import datetime
from io import BytesIO
from docx import Document
import pandas as pd

# ---------- PROJECT ROOT ----------

project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

# ---------- PAGE CONFIG ----------

st.set_page_config(
    page_title="IJ Global Scraper",
    page_icon="🌐",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ---------- CSS ----------

st.markdown("""
<style>

#MainMenu {visibility:hidden;}
footer {visibility:hidden;}

/* Animated progress bar */

.stProgress > div > div > div > div {
background: linear-gradient(90deg,#4facfe,#00f2fe,#4facfe);
background-size: 200% 100%;
animation: progressAnimation 2s linear infinite;
}

@keyframes progressAnimation {
0% {background-position:0%}
100% {background-position:200%}
}
.loader {
border:6px solid #f3f3f3;
border-top:6px solid #1f77b4;
border-radius:50%;
width:50px;
height:50px;
animation:spin 1s linear infinite;
margin:auto;
}

@keyframes spin{
0%{transform:rotate(0deg);}
100%{transform:rotate(360deg);}
}

/* Header styling with flexbox */
.header-wrapper {
display: flex;
align-items: center;
justify-content: space-between;
gap: 20px;
margin-bottom: 1.5rem;
padding: 10px 0;
}

.header-left {
display: flex;
align-items: center;
gap: 20px;
flex: 1;
min-width: 0;
}

.header-logo {
flex-shrink: 0;
width: 180px;
height: auto;
}

.header-logo img {
width: 100%;
height: auto;
display: block;
}

.header-text {
flex: 1;
min-width: 0;
}

.main-header {
font-size: 2.6rem;
font-weight: 700;
color: #1f77b4;
margin: 0;
line-height: 1.2;
word-wrap: break-word;
}

.sub-header {
font-size: 1.1rem;
color: #666;
margin: 0.25rem 0 0 0;
line-height: 1.4;
word-wrap: break-word;
}

.header-button {
flex-shrink: 0;
}

.preview-btn button {
padding: 6px 10px;
font-size: 14px;
}

.tab-button {
padding: 10px 20px;
font-size: 16px;
font-weight: 600;
}
/* Active primary button color */
.stButton button[kind="primary"] {
background-color: #1f77b4;
border-color: #1f77b4;
color: white;
}

.stButton button[kind="primary"]:hover {
background-color: #16639a;
border-color: #16639a;
}
            
section[data-testid="stSidebar"] .stMarkdown {
    margin-bottom: -20px;
}

section[data-testid="stSidebar"] .stSelectbox {
    margin-top: -8px;
}

section[data-testid="stSidebar"] .stMultiSelect {
    margin-top: -8px;
}

/* Hide dialog close (X) button */
[data-testid="stBaseButton-headerNoPadding"] {
    display: none !important;
}

/* Deploy and Stop button font color */
[data-testid="stToolbar"] button,
[data-testid="stToolbar"] a,
.stDeployButton {
    color: white !important;
}

[data-testid="stToolbar"] button:hover,
[data-testid="stToolbar"] a:hover,
.stDeployButton:hover {
    color: white !important;
    background-color: white !important;
    border-color: white !important;
}

</style>
""", unsafe_allow_html=True)

# ---------- HEADER ----------

# Initialize session state for filter visibility
if 'show_filters' not in st.session_state:
    st.session_state.show_filters = False

if "open_dialog" not in st.session_state:
    st.session_state.open_dialog = False

if 'current_tab' not in st.session_state:
    st.session_state.current_tab = "Run Scraper"

if 'session_reports' not in st.session_state:
    st.session_state.session_reports = []

if 'scraper_config' not in st.session_state:
    st.session_state.scraper_config = {
        'region': None,
        'country': None,
        'industry_type': None,
        'website': None
    }

if 'scraper_config' not in st.session_state:
    st.session_state.scraper_config = {
        'region': None,
        'country': None,
        'industry_type': None,
        'website': None
    }

# Modal popup dialog for filters
@st.dialog("🔍 Relevance Phrases & Keywords")
def show_filters_dialog():
    # Load relevance filters from JSON
    relevance_config_path = os.path.join(project_root, "config", "relevantnews.json")
    with open(relevance_config_path) as f:
        relevance_data = json.load(f)
    
    phrases = relevance_data["relevanceFilters"]["phrases"]
    keywords = relevance_data["relevanceFilters"]["keywords"]
    
    # Display in two columns
    filter_col1, filter_col2 = st.columns(2)
    
    with filter_col1:
        st.subheader("📋 Relevance Phrases")
        
        # Display existing phrases in a clean list
        if phrases:
            # Create a selectbox for choosing which phrase to edit/delete
            selected_phrase = st.selectbox(
                "Select phrase to edit/delete:",
                ["None"] + phrases,
                key="phrase_selector"
            )
            
            if selected_phrase != "None":
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("✏️ Edit Selected", key="edit_selected_phrase", use_container_width=True):
                       st.session_state.editing_phrase = selected_phrase
                       st.session_state.open_dialog = True
                       st.rerun()
                with col2:
                    if st.button("🗑️ Delete Selected", key="delete_selected_phrase", use_container_width=True):
                        phrases.remove(selected_phrase)
                        relevance_data["relevanceFilters"]["phrases"] = phrases
                        with open(relevance_config_path, "w") as f:
                            json.dump(relevance_data, f, indent=2)
                        st.success(f"✅ Deleted: {selected_phrase}")
                        st.rerun()
            
            # Edit mode
            if st.session_state.get("editing_phrase"):
                st.markdown("**Edit Phrase:**")
                edited_phrase = st.text_input(
                    "Edit phrase:", 
                    value=st.session_state.editing_phrase, 
                    key="edit_phrase_input"
                )
                col_save, col_cancel = st.columns(2)
                with col_save:
                    if st.button("💾 Save", key="save_phrase", use_container_width=True):
                        if edited_phrase and edited_phrase.strip():
                            phrase_index = phrases.index(st.session_state.editing_phrase)
                            phrases[phrase_index] = edited_phrase.strip()
                            relevance_data["relevanceFilters"]["phrases"] = phrases
                            with open(relevance_config_path, "w") as f:
                                json.dump(relevance_data, f, indent=2)
                            del st.session_state.editing_phrase
                            st.success(f"✅ Updated: {edited_phrase.strip()}")
                            st.rerun()
                with col_cancel:
                    if st.button("❌ Cancel", key="cancel_phrase", use_container_width=True):
                        del st.session_state.editing_phrase
                        st.rerun()
            
            st.markdown("---")
            st.markdown("**Current Phrases:**")
            for phrase in phrases:
                st.markdown(f"• {phrase}")
        else:
            st.info("No phrases added yet.")
        
        st.markdown("---")
        st.markdown("**Add New Phrase**")
        new_phrase = st.text_input("📝 Enter new phrase", key="new_phrase_input", placeholder="Type phrase here...")
        if st.button("➕ Add Phrase", key="add_phrase_btn", use_container_width=True):
            if new_phrase and new_phrase.strip():
                if new_phrase.strip() not in phrases:
                    phrases.append(new_phrase.strip())
                    relevance_data["relevanceFilters"]["phrases"] = phrases
                    with open(relevance_config_path, "w") as f:
                        json.dump(relevance_data, f, indent=2)
                    st.success(f"✅ Added: {new_phrase.strip()}")
                    st.rerun()
                else:
                    st.warning("⚠️ Phrase already exists!")
            else:
                st.error("⚠️ Please enter a valid phrase!")
    
    with filter_col2:
        st.subheader("🔑 Relevance Keywords")
        
        # Display existing keywords in a clean list
        if keywords:
            # Create a selectbox for choosing which keyword to edit/delete
            selected_keyword = st.selectbox(
                "Select keyword to edit/delete:",
                ["None"] + keywords,
                key="keyword_selector"
            )
            
            if selected_keyword != "None":
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("✏️ Edit Selected", key="edit_selected_keyword", use_container_width=True):
                       st.session_state.editing_keyword = selected_keyword
                       st.session_state.open_dialog = True
                       st.rerun()
                with col2:
                    if st.button("🗑️ Delete Selected", key="delete_selected_keyword", use_container_width=True):
                        keywords.remove(selected_keyword)
                        relevance_data["relevanceFilters"]["keywords"] = keywords
                        with open(relevance_config_path, "w") as f:
                            json.dump(relevance_data, f, indent=2)
                        st.success(f"✅ Deleted: {selected_keyword}")
                        st.rerun()
            
            # Edit mode
            if st.session_state.get("editing_keyword"):
                st.markdown("**Edit Keyword:**")
                edited_keyword = st.text_input(
                    "Edit keyword:", 
                    value=st.session_state.editing_keyword, 
                    key="edit_keyword_input"
                )
                col_save, col_cancel = st.columns(2)
                with col_save:
                    if st.button("💾 Save", key="save_keyword", use_container_width=True):
                        if edited_keyword and edited_keyword.strip():
                            keyword_index = keywords.index(st.session_state.editing_keyword)
                            keywords[keyword_index] = edited_keyword.strip()
                            relevance_data["relevanceFilters"]["keywords"] = keywords
                            with open(relevance_config_path, "w") as f:
                                json.dump(relevance_data, f, indent=2)
                            del st.session_state.editing_keyword
                            st.success(f"✅ Updated: {edited_keyword.strip()}")
                            st.rerun()
                with col_cancel:
                    if st.button("❌ Cancel", key="cancel_keyword", use_container_width=True):
                        del st.session_state.editing_keyword
                        st.rerun()
            
            st.markdown("---")
            st.markdown("**Current Keywords:**")
            for keyword in keywords:
                st.markdown(f"• {keyword}")
        else:
            st.info("No keywords added yet.")
        
        st.markdown("---")
        st.markdown("**Add New Keyword**")
        new_keyword = st.text_input("📝 Enter new keyword", key="new_keyword_input", placeholder="Type keyword here...")
        if st.button("➕ Add Keyword", key="add_keyword_btn", use_container_width=True):
            if new_keyword and new_keyword.strip():
                if new_keyword.strip() not in keywords:
                    keywords.append(new_keyword.strip())
                    relevance_data["relevanceFilters"]["keywords"] = keywords
                    with open(relevance_config_path, "w") as f:
                        json.dump(relevance_data, f, indent=2)
                    st.success(f"✅ Added: {new_keyword.strip()}")
                    st.rerun()
                else:
                    st.warning("⚠️ Keyword already exists!")
            else:
                st.error("⚠️ Please enter a valid keyword!")
    
    st.markdown("---")
    
    # Close button
    if st.button("❌ Close", use_container_width=True):

       st.session_state.open_dialog = False

       if "editing_phrase" in st.session_state:
           del st.session_state.editing_phrase

       if "editing_keyword" in st.session_state:
           del st.session_state.editing_keyword

       st.rerun()

# Header
st.markdown('<div class="main-header">🌐 Scraper & Report Viewer</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Configure and run web scraping tasks, then view generated reports</div>', unsafe_allow_html=True)

# Phrases button in sidebar area
with st.columns([1])[0]:
    if st.button("🔑 Phrases/Keywords", key="filters_toggle", use_container_width=True):
        st.session_state.open_dialog = True

if st.session_state.open_dialog:
    show_filters_dialog()
    st.session_state.open_dialog = False  

# ---------- NAVIGATION TABS ----------

tab_col1, tab_col2 = st.columns(2)

# Determine button types based on active tab
run_scraper_type = "primary" if st.session_state.current_tab == "Run Scraper" else "secondary"
past_reports_type = "primary" if st.session_state.current_tab == "Past Reports" else "secondary"

with tab_col1:
    if st.button("▶️ Generate Report", use_container_width=True, key="tab_run_scraper", type=run_scraper_type):
        st.session_state.current_tab = "Run Scraper"
        st.rerun()

with tab_col2:
    if st.button("📚 Past Reports", use_container_width=True, key="tab_past_reports", type=past_reports_type):
        st.session_state.current_tab = "Past Reports"
        st.rerun()

st.markdown("---")

# ---------- PREVIEW POPUP ----------




# ---------- PREVIEW POPUP ----------

@st.dialog("📄 Report Preview")
def show_preview(file_path):

    doc = Document(file_path)

    paragraphs = []
    tables_data = []

    # -------- Read paragraphs --------
    for p in doc.paragraphs:
        if p.text.strip():
            paragraphs.append(p.text)

    # -------- Read tables --------
    for table in doc.tables:
        table_rows = []

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())

            table_rows.append(row_data)

        tables_data.append(table_rows)

    # -------- Show paragraph content --------
    edited_text = st.text_area(
        "Report Content",
        "\n".join(paragraphs),
        height=250
    )

    # -------- Show tables with borders --------
    for table in tables_data:

        st.markdown("### Parameters Table")

        df = pd.DataFrame(table[1:], columns=table[0])

        st.table(df)

    # -------- Buttons --------
    col1, col2 = st.columns(2)

    with col1:
        if st.button("💾 Save Changes"):
            st.success("Edit saved locally (demo)")

    with col2:
        if st.button("❌ Close"):
            st.rerun()

# ---------- METADATA POPUP ----------

@st.dialog("📋 Report Metadata")
def show_metadata(file_path):
    
    doc = Document(file_path)
    
    # Extract metadata from document
    file_name = os.path.basename(file_path)
    mod_time = datetime.fromtimestamp(os.path.getmtime(file_path))
    
    # Try to load metadata from JSON file
    json_file_path = file_path.replace(".docx", ".json")
    metadata_loaded = False
    
    if os.path.exists(json_file_path):
        try:
            with open(json_file_path, "r") as json_file:
                metadata = json.load(json_file)
                metadata_loaded = True
        except:
            metadata = {}
    else:
        metadata = {}
    
    st.subheader("📋 Report Information")
    
    st.write(f"**File Name:** {file_name}")
    st.write(f"**Generated Date:** {metadata.get('generated_date', mod_time.strftime('%Y-%m-%d %H:%M:%S'))}")
    st.write(f"**File Size:** {metadata.get('file_size_kb', os.path.getsize(file_path) / 1024):.1f} KB")
    
    st.markdown("---")
    
    # Get metadata from JSON or fallback to session state or document
    if metadata_loaded:
        project_name = metadata.get("project_name", "N/A")
        country = metadata.get("country", "N/A")
        region = metadata.get("region", "N/A")
        industry_type = metadata.get("industry_type", "N/A")
        website = metadata.get("website", "N/A")
        source_url = metadata.get("source_url", "N/A")
    else:
        # Fallback to session state and document extraction
        project_name = "N/A"
        country = st.session_state.scraper_config.get('country', 'N/A') or 'N/A'
        region = st.session_state.scraper_config.get('region', 'N/A') or 'N/A'
        industry_type = st.session_state.scraper_config.get('industry_type', 'N/A') or 'N/A'
        website = st.session_state.scraper_config.get('website', 'N/A') or 'N/A'
        source_url = "N/A"
        
        # Get project name from first paragraph
        if len(doc.paragraphs) > 0:
            project_name = doc.paragraphs[0].text if doc.paragraphs[0].text else "N/A"
        
        # Extract additional data from tables if available
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    key = row.cells[0].text.strip()
                    value = row.cells[1].text.strip()
                    
                    if key.lower() == "country" and value and value != "Not Found":
                        country = value
                    elif key.lower() == "project name" and value and value != "Not Found":
                        project_name = value
    
    st.subheader("📊 Project Details")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.write(f"**Project Name:** {project_name}")
        st.write(f"**Country:** {country}")
        if metadata_loaded:
            st.write(f"**Website:** {website}")
    
    with col2:
        st.write(f"**Region:** {region}")
        st.write(f"**Industry Type:** {industry_type}")
        if metadata_loaded and source_url != "N/A":
            st.write(f"**Source URL:** {source_url[:50]}...")
    
    st.markdown("---")
    
    if st.button("❌ Close", use_container_width=True):

        st.session_state.open_dialog = False   

        if "editing_phrase" in st.session_state:
            del st.session_state.editing_phrase

        if "editing_keyword" in st.session_state:
            del st.session_state.editing_keyword

        st.rerun()

# ---------- LOAD CONFIG ----------

config_path = os.path.join(project_root, "config", "sites.json")

with open(config_path) as f:
    data = json.load(f)

sites = data["sites"]

# ---------- TAB 1: RUN SCRAPER ----------

if st.session_state.current_tab == "Run Scraper":
    
    # ---------- SIDEBAR ----------
    
    with st.sidebar:
        
        # Display logo at top of sidebar
        logo_path = os.path.join(project_root, "ui", "static", "ijglobal_logo.png")
        if os.path.exists(logo_path):
            st.image(logo_path, width=280)
        
        st.header("⚙️ Configuration")
    
        regions = sorted(list(set([s["region"] for s in sites])))
    
        st.markdown("📍 Region <span style='color:red'>*</span>", unsafe_allow_html=True)
        

        region = st.selectbox(
             "",
             regions,
             index=None,
             placeholder="Select Region"
        )
        
        # Store in session state
        st.session_state.scraper_config['region'] = region
        
        # Store in session state
        st.session_state.scraper_config['region'] = region
    
        if region:
            countries = sorted(list(set([s["country"] for s in sites if s["region"] == region])))
        else:
            countries = []
    
        st.markdown("🌍 Country <span style='color:red'>*</span>", unsafe_allow_html=True)

        country = st.selectbox(
              "",
              countries,
              index=None,
              placeholder="Select Country"
        )
        
        # Store in session state
        st.session_state.scraper_config['country'] = country
    
        if country:
            filtered_sites = [s for s in sites if s["country"] == country]
            website_names = [s["siteName"] for s in filtered_sites]
        else:
            filtered_sites = []
            website_names = []
    
        st.markdown("🔗 Website <span style='color:red'>*</span>", unsafe_allow_html=True)

        website = st.selectbox(
              "",
              website_names,
              index=None,
              placeholder="Select Website"
        )
        
        # Store in session state
        st.session_state.scraper_config['website'] = website
    
        selected_site = None
    
        if website:
            selected_site = next((s for s in filtered_sites if s["siteName"] == website), None)
    
        st.divider()
    
        if selected_site:
            industry_types = selected_site.get("industryType", [])
            if isinstance(industry_types, str):
                industry_types = [industry_types]
        else:
            industry_types = []
    
        st.markdown("🏭 Industry Type <span style='color:red'>*</span>", unsafe_allow_html=True)

        selected_industries = st.multiselect(
            "",
            industry_types
        )
        
        # Store in session state
        st.session_state.scraper_config['industry_type'] = ', '.join(selected_industries) if selected_industries else None
        
        # Store in session state
        st.session_state.scraper_config['industry_type'] = ', '.join(selected_industries) if selected_industries else None
    
        days = st.number_input(
            "📅 Number of Days",
            min_value=1,
            max_value=365,
            value=7
        )
    
        st.divider()
    
        is_disabled = not (region and country and website and selected_industries)

        run_button = st.button(
           "▶️ Generate Report",
           disabled=is_disabled
        )
    
    # ---------- MAIN AREA ----------
    
    col1, col2 = st.columns([2,1])
    
    with col1:
    
        st.subheader("📊 Scraper Configuration")
    
        c1, c2 = st.columns(2)
    
        with c1:
    
            st.info(f"**Region:** {region if region else 'Not Selected'}")
            st.info(f"**Country:** {country if country else 'Not Selected'}")
            st.info(f"**Website:** {website if website else 'Not Selected'}")
    
        with c2:
    
            st.info(f"**Industries:** {', '.join(selected_industries) if selected_industries else 'None'}")
            st.info(f"**Days Range:** {days}")
    
            if selected_site:
                st.info(f"**URL:** {selected_site['siteURL']}")
    
    with col2:
    
        st.subheader("📈 Statistics")
        st.metric("Processed Records", len(st.session_state.session_reports))
    
    # ---------- RUN SCRAPER ----------
    
    if run_button:
    
        if not region or not country or not website:
    
            st.error("⚠️ Please select Region, Country and Website")
    
        elif not selected_industries:
    
            st.error("⚠️ Please select Industry")
    
        else:
    
            progress_bar = st.progress(0)
    
            step_indicator = st.empty()
            record_counter = st.empty()
    
            loader = st.empty()
            loader.markdown("""
<div style="text-align:center">
<div class="loader"></div>
<h3>In progress...</h3>
</div>
""", unsafe_allow_html=True)
    
            step_indicator.info("Step 1/4 : Initializing scraper")

            record_counter.metric("Processed Records", 0)
 
            step_indicator.info("Step 1/4 : Initializing scraper")
            progress_bar.progress(10)
            time.sleep(0.5)
 
            step_indicator.info("Step 2/4 : Opening website")
            progress_bar.progress(30)
            time.sleep(0.5)
 
            step_indicator.info("Step 3/4 : Extracting project data")
            progress_bar.progress(60)
   
            #onedrive_path = os.getenv("OneDriveCommercial") or os.getenv("OneDrive")
            #output_dir = os.path.join(onedrive_path,"IJ Global Extracted File")
            step_indicator.info("Step 4/4 : Generating reports")
            progress_bar.progress(90)
            time.sleep(0.5)
            output_dir = os.path.join(project_root, "output")
            os.makedirs(output_dir, exist_ok=True)
    
    
            # initial file count
            initial_files = 0
            if os.path.exists(output_dir):
                initial_files = len(glob.glob(os.path.join(output_dir,"*.docx")))

            run_start_time = time.time()
    
            # Check if Korean website is selected
            if website == "Korea Dart":
                import Testkorean
                with st.spinner("Running Korean scraper..."):
                    Testkorean.run()
            else:
                import subprocess
                scraper_script = os.path.join(project_root, "scripts", "scrapper.py")
                with st.spinner("Running scraper..."):
                    result = subprocess.run(
                        [sys.executable, scraper_script],
                        capture_output=True, text=True, cwd=project_root
                    )
                if result.returncode != 0:
                    st.error(f"Scraper error:\n{result.stderr}")
                    st.stop()

            progress_bar.progress(1.0)

            loader.empty()

            step_indicator.success("✅ Scraping Completed")

            st.balloons()

            # Update records after completion
            current_files = 0
            if os.path.exists(output_dir):
                current_files = len(glob.glob(os.path.join(output_dir,"*.docx")))

            new_records = current_files - initial_files

            record_counter.metric("Processed Records", new_records)
            
            # Store session reports
            if os.path.exists(output_dir):
                all_files = glob.glob(os.path.join(output_dir, "*.docx"))
                st.session_state.session_reports = sorted(
                    [f for f in all_files if os.path.getmtime(f) >= run_start_time],
                    key=os.path.getmtime, reverse=True
                )
    
    st.divider()
    
    # ---------- SESSION REPORTS ----------
    
    st.subheader("📁 Reports from Current Session")
    
    if st.session_state.session_reports:
        
        session_reports = sorted(st.session_state.session_reports, key=os.path.getmtime, reverse=True)
        
        for report in session_reports:
    
            name = os.path.basename(report)
            size = os.path.getsize(report)/1024
            time_mod = datetime.fromtimestamp(os.path.getmtime(report))
    
            c1,c2,c3,c4,c5,c6 = st.columns([2.5,2,1,0.5,0.5,0.5])
    
            with c1:
                st.write(f"📄 {name}")
    
            with c2:
                st.write(time_mod.strftime("%Y-%m-%d %H:%M"))
    
            with c3:
                st.write(f"{size:.1f} KB")
    
            with c4:
                if st.button("👁", key=f"preview_{name}"):
                    show_preview(report)
    
            with c5:
                if st.button("ℹ️", key=f"meta_{name}"):
                    show_metadata(report)
    
            with c6:
                with open(report,"rb") as file:
                    st.download_button(
                        "⬇️",
                        file,
                        file_name=name,
                        key=f"dl_{name}"
                    )
    
            st.divider()
    
    else:
        st.info("📭 No reports generated in current session. Run the scraper to generate reports.")

# ---------- TAB 2: PAST REPORTS ----------

elif st.session_state.current_tab == "Past Reports":
    
    # ---------- SIDEBAR FOR PAST REPORTS ----------
    
    with st.sidebar:
        
        # Display logo at top of sidebar
        logo_path = os.path.join(project_root, "ui", "static", "ijglobal_logo.png")
        if os.path.exists(logo_path):
            st.image(logo_path, width=280)
        
        st.header("🔎 Filter Reports")
        
        #onedrive_path = os.getenv("OneDriveCommercial") or os.getenv("OneDrive")

        output_dir = os.path.join(project_root, "output")
        os.makedirs(output_dir, exist_ok=True)

        if os.path.exists(output_dir):

            all_reports = glob.glob(os.path.join(output_dir, "*.docx"))

            # Load all metadata
            all_metadata = []
            for report in all_reports:
                json_file = report.replace(".docx", ".json")
                if os.path.exists(json_file):
                    try:
                        with open(json_file, "r") as f:
                            metadata = json.load(f)
                            all_metadata.append(metadata)
                    except:
                        pass

            # Extract unique values for filters
            countries = sorted(list(set([m.get("country", "N/A") for m in all_metadata if m.get("country")])))
            regions = sorted(list(set([m.get("region", "N/A") for m in all_metadata if m.get("region")])))
            industries = sorted(list(set([m.get("industry_type", "N/A") for m in all_metadata if m.get("industry_type")])))
            websites = sorted(list(set([m.get("website", "N/A") for m in all_metadata if m.get("website")])))

            # Search by filename
            search_term = st.text_input("📄 Search by Filename", placeholder="Enter File Name")

            # Country filter
            selected_country = st.selectbox(
                "🌍 Country",
                ["All Countries"] + countries,
                index=0
            )

            # Region filter
            selected_region = st.selectbox(
                "📍 Region",
                ["All Regions"] + regions,
                index=0
            )

            # Industry Type filter
            selected_industry = st.selectbox(
                "🏭 Industry Type",
                ["All Industries"] + industries,
                index=0
            )

            # Website filter
            selected_website = st.selectbox(
                "🔗 Website",
                ["All Websites"] + websites,
                index=0
            )

            # Date filter
            date_filter = st.selectbox(
                "📅 Generated Date",
                ["All Dates", "Today", "Last 7 Days", "Last 30 Days"],
                index=0
            )

        else:
            search_term = ""
            selected_country = "All Countries"
            selected_region = "All Regions"
            selected_industry = "All Industries"
            selected_website = "All Websites"
            date_filter = "All Dates"
    
    # ---------- MAIN AREA - ALL REPORTS ---------- 

    st.subheader("📚 All Generated Reports")

    output_dir = os.path.join(project_root, "output")

    if os.path.exists(output_dir):

        all_reports = glob.glob(os.path.join(output_dir, "*.docx"))

        if all_reports:

            # Load metadata for all reports
            reports_with_metadata = []
            for report in all_reports:
                json_file = report.replace(".docx", ".json")
                metadata = {}

                if os.path.exists(json_file):
                    try:
                        with open(json_file, "r") as f:
                            metadata = json.load(f)
                    except:
                        pass

                reports_with_metadata.append({
                    "path": report,
                    "metadata": metadata
                })

            # Apply filters
            filtered_reports = reports_with_metadata

            if search_term:
                filtered_reports = [
                    r for r in filtered_reports 
                    if search_term.lower() in os.path.basename(r["path"]).lower()
                ]

            if selected_country != "All Countries":
                filtered_reports = [
                    r for r in filtered_reports 
                    if r["metadata"].get("country") == selected_country
                ]

            if selected_region != "All Regions":
                filtered_reports = [
                    r for r in filtered_reports 
                    if r["metadata"].get("region") == selected_region
                ]

            if selected_industry != "All Industries":
                filtered_reports = [
                    r for r in filtered_reports 
                    if r["metadata"].get("industry_type") == selected_industry
                ]

            if selected_website != "All Websites":
                filtered_reports = [
                    r for r in filtered_reports 
                    if r["metadata"].get("website") == selected_website
                ]

            if date_filter != "All Dates":
                today = datetime.now().date()
                filtered_by_date = []

                for r in filtered_reports:
                    gen_date_str = r["metadata"].get("generated_date", "")

                    if gen_date_str:
                        try:
                            gen_date = datetime.strptime(gen_date_str, "%Y-%m-%d %H:%M:%S").date()

                            if date_filter == "Today" and gen_date == today:
                                filtered_by_date.append(r)
                            elif date_filter == "Last 7 Days" and (today - gen_date).days <= 7:
                                filtered_by_date.append(r)
                            elif date_filter == "Last 30 Days" and (today - gen_date).days <= 30:
                                filtered_by_date.append(r)
                        except:
                            pass

                filtered_reports = filtered_by_date

            filtered_reports.sort(
                key=lambda x: os.path.getmtime(x["path"]), 
                reverse=True
            )

            st.write(f"Showing {len(filtered_reports)} of {len(all_reports)} reports")

            if filtered_reports:
                zip_buffer = BytesIO()

                with zipfile.ZipFile(zip_buffer, "w") as z:
                    for r in filtered_reports:
                        z.write(r["path"], os.path.basename(r["path"]))

                st.download_button(
                    "⬇️ Download All Filtered Reports",
                    zip_buffer.getvalue(),
                    file_name="IJ_Global_Reports.zip"
                )

            st.divider()

            for report_item in filtered_reports[:50]:

                report = report_item["path"]
                metadata = report_item["metadata"]

                name = os.path.basename(report)
                size = os.path.getsize(report) / 1024
                time_mod = datetime.fromtimestamp(os.path.getmtime(report))

                c1, c2, c3, c4, c5, c6 = st.columns([2.5, 2, 1, 0.5, 0.5, 0.5])

                with c1:
                    st.write(f"📄 {name}")

                with c2:
                    st.write(time_mod.strftime("%Y-%m-%d %H:%M"))

                with c3:
                    st.write(f"{size:.1f} KB")

                with c4:
                    if st.button("👁", key=f"preview_past_{name}"):
                        show_preview(report)

                with c5:
                    if st.button("ℹ️", key=f"meta_past_{name}"):
                        show_metadata(report)

                with c6:
                    with open(report, "rb") as file:
                        st.download_button(
                            "⬇️",
                            file,
                            file_name=name,
                            key=f"dl_past_{name}"
                        )

                st.divider()

            if len(filtered_reports) > 50:
                st.info(f"ℹ️ Showing first 50 reports. Total: {len(filtered_reports)}")

        else:
            st.info("📭 No reports found")

    else:
        st.warning("⚠️ Output folder not found")
